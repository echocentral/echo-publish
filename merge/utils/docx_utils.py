#!/usr/bin/python
# -*- coding: iso-8859-15 -*-
import os
import io
import re
import string
import zipfile
import tempfile
import shutil
import xmltodict
from merge.utils.engine_utils import substituteVariablesPlainString
import django.template.exceptions
from docx import Document
"""
from django.conf import settings
#  from docx.text.paragraph import Paragraph
from django.template import Context, Engine
from markdown import markdown
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from subprocess import check_output
import lxml.etree as etree
import PyPDF2
from pynliner import Pynliner

from .resource_utils import (strip_xml_dec, get_xml_dec)
from .config import local_root
"""

"""
def docx_copy_run_style_from(run1, run2):
    run1.font.color.rgb = run2.font.color.rgb
    run1.font.all_caps = run2.font.all_caps
    run1.font.bold = run2.font.bold
    run1.font.italic = run2.font.italic
    run1.font.size = run2.font.size
    run1.font.underline = run2.font.underline
    # complex_script, cs_bold, cs_italic, double_strike, emboss, hidden, highlight_color,
    # imprint, math, name, no_proof, outline, rtl, shadow, small_caps, snap_to_grid, spec_vanish, 
    # strike, superscript, underline, web_hidden


def docx_copy_para_format_from(para1, para2):
    para1.paragraph_format.alignment = para2.paragraph_format.alignment
    para1.paragraph_format.first_line_indent = para2.paragraph_format.first_line_indent
    para1.paragraph_format.keep_together = para2.paragraph_format.keep_together
    para1.paragraph_format.keep_with_next = para2.paragraph_format.keep_with_next
    para1.paragraph_format.left_indent = para2.paragraph_format.left_indent
    try:
        para1.paragraph_format.line_spacing = para2.paragraph_format.line_spacing
        para1.paragraph_format.line_spacing_rule = para2.paragraph_format.line_spacing_rule
    except ValueError:
        pass
    para1.paragraph_format.page_break_before = para2.paragraph_format.page_break_before
    para1.paragraph_format.right_indent = para2.paragraph_format.right_indent
    para1.paragraph_format.space_after = para2.paragraph_format.space_after
    para1.paragraph_format.space_before = para2.paragraph_format.space_before
    para1.paragraph_format.widow_control = para2.paragraph_format.widow_control

"""
def isControlText(s):
    s = s.strip()
    if s[:2] == "{%" and s[-2:] == "%}" and s.find("%}") == s.rfind("%}"):
        if s.find("include") >= 0:
            return False
        elif s.find("now") >= 0:
            return False
        else:
            return True
    else:
        return False

def removePara(para):
    p = para._element
    p.getparent().remove(p)
    p._p = p._element = None
        
"""

def isControlLine(s):
    s = s.split("+")[0]
    return isControlText(s)


def wrap_list_xml(childlist, root_tag, child_tag):
    out = "<"+root_tag+">\n"
    for child in childlist:
        out += "\t<"+child_tag+">"+child+"</"+child_tag+">\n"
    out += "</"+root_tag+">"
    return out


def table_text(table):
    fullText = "" 
    rows = table.rows
    for row in rows:
        cells = row.cells
        for cell in cells:
            paras = cell.paragraphs
            for para in paras:
                paraText = para.text
                fullText += paraText
            subtables = cell.tables
            for subtable in subtables:
                fullText += table_text(subtable)
    return fullText

def docx_text(file_name_in):
    doc_in = Document(docx=file_name_in)
    paras = doc_in.paragraphs
    fullText = "" 
    for para in paras:
        paraText = para.text
        fullText += paraText
    tables = doc_in.tables
    for table in tables:
        fullText += table_text(table)
    return fullText


def extract_regex_matches_docx(file_name_in, regex, wrap=None, root_tag="list", child_tag="item"):
    text = docx_text(file_name_in)
    p = re.compile(regex)
    m = p.findall(text)
    if wrap:
        return wrap_list_xml(m, root_tag, child_tag)
    else:
        return m

"""
def preprocess_docx_template(file_name_in, file_name_out):
    timestamp_template = os.stat(file_name_in).st_mtime
    prep_file_exists = False
    if os.path.isfile(file_name_out):
        timestamp_template_prep = os.stat(file_name_out).st_mtime
        prep_file_exists = True
    if (not prep_file_exists) or timestamp_template_prep < timestamp_template:
        doc_in = Document(docx=file_name_in)
        paras = doc_in.paragraphs
        for para in paras:
            runs = para.runs
            for run in runs:
                txt = run.text
                if isControlText(txt):
                    run.text = "[##]"+txt
        doc_in.save(file_name_out)
"""
"""
def process_newlines(para):
    runs = para.runs
    for run in runs:
        txt = run.text
        if (txt.find("\\n")) >= 0:
            run.text = ""
            while (txt.find("\\n")) >= 0:
                run.add_text(txt[:txt.find("\\n")])
                run.add_break()
                txt = txt[txt.find("\\n")+2:]
            run.add_text(txt)
"""
"""
def postprocess_docx(file_name_in):
    doc_in = Document(docx=file_name_in)
    for p in doc_in.paragraphs:
        if p.text.strip() == "[##]":
            removePara(p)
        if p.text.find("[##]") >= 0:
            p.text = p.text.replace("[##]","")
        process_newlines(p)
    for t in doc_in.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    process_newlines(p)

    doc_in.save(file_name_in)
"""

def substituteVariablesDocx(config, file_name_in, fileNameOut, subs):
    c = Context(subs)
    doc_in = Document(docx=file_name_in.replace("/./", "/").replace("\\.\\", "\\").replace("\\", "/"))
    doc_temp = Document()
    paras = doc_in.paragraphs
    fullText = "" 
    i = 0
    for para in paras:
        paraText = ""
        p = doc_temp.add_paragraph(style=para.style)
        docx_copy_para_format_from(p, para)
        j = 0
        runs = para.runs
        for run in runs:
            txt = run.text
            paraText += txt+"+"+str(j)+"+run+"
            r = p.add_run(text=txt, style=run.style)
            docx_copy_run_style_from(r, run)
            j += 1
        fullText += paraText+str(i)+"+para+"
        i += 1
    fullText = preprocess(fullText)
    t = get_engine(config).from_string(fullText)
    xtxt = t.render(c)
    xtxt = apply_sequence(xtxt)
    xParaTxts = xtxt.split("+para+")
    for p in paras:
        removePara(p)

    doc_in.paragraphs.clear()
    paras = doc_temp.paragraphs
    for xParaTxt in xParaTxts:
        runTxts = xParaTxt.split("+run+")
        if runTxts[-1] != '':
            para_n = int(runTxts[-1])
            p = doc_in.add_paragraph(style=paras[para_n].style)
            docx_copy_para_format_from(p, paras[para_n])
            for runTxt in runTxts[:-1]:
                try:
                    txt = runTxt.split("+")[-2]
                except:
                    txt = ""
                run_n = int(runTxt.split("+")[-1])
                r = p.add_run(text=txt, style=paras[para_n].runs[run_n].style)
                docx_copy_run_style_from(r, paras[para_n].runs[run_n])
            if isControlLine(paras[para_n].text):
                p.text = "{}"

    for p in doc_in.paragraphs:
        if p.text == "{}":
            removePara(p)
    doc_in.save(fileNameOut)
    return {"file": fileNameOut}


def print_doc(doc):
    paras = doc.paragraphs
    for para in paras[14:20]:
        print(para.text)


def combine_docx(file_names, file_name_out):
    combined_document = Document(file_names[0])
    for file in file_names[1:]:
        if file == "pagebreak":
            combined_document.add_page_break()
        else:    
            sub_doc = Document(file)
            for para in sub_doc.paragraphs:
                pnew = combined_document.add_paragraph(style=para.style)
                docx_copy_para_format_from(pnew, para)
                runs = para.runs
                for run in runs:
                    rnew = pnew.add_run(text=run.text, style=run.style)
                    docx_copy_run_style_from(rnew, run)

    combined_document.save(file_name_out)
    return {"file": file_name_out}


def convert_markdown(fileNameIn, fileNameOut, css_string=None):
    fileIn = open(fileNameIn, "r", encoding="utf-8")
    fileOut = open(fileNameOut, "w", encoding="utf-8")
    html = markdown(fileIn.read())
    if css_string:
        html = '<div class="echo-publish">'+html+"</div>"
        html = Pynliner().from_string(html).with_cssString(css_string).run()
    fileOut.write(html)
    fileOut.close()
    return {"file": fileNameOut}


def convert_markdown_string(stringIn):
    return markdown(stringIn, extensions=['markdown.extensions.admonition','markdown.extensions.tables','markdown.extensions.nl2br'])


def email_file(baseFileName, me, you, subject, credentials):

    msg = MIMEMultipart('alternative')
    msg['Subject'] = subject
    msg['From'] = me
    msg['To'] = you.replace(" ","+")

    mimetypes = ["text/plain", "text/html"]
    file_exts = [".md", ".html"]

    for mime in zip(mimetypes, file_exts):
        file=baseFileName+mime[1]
        fp = open(file, 'r')
        content = fp.read()
        fp.close()
        msg.attach(MIMEText(content, mime[0].split("/")[1]))    
    
    username = credentials["username"]
    password = credentials["password"]
    server = smtplib.SMTP(credentials["server"])

    server.ehlo()
    server.starttls()
    server.login(username, password)
    server.sendmail(me, [you.replace(" ", "+")], msg.as_string())
    server.quit()
    return {"email": you.replace(" ", "+")}


def merge_docx_footer(config, full_local_filename, subs):
    merge_docx_header_footer(config, full_local_filename, subs, "footer1")
    return merge_docx_header_footer(config, full_local_filename, subs, "footer2")


def merge_docx_header(config, full_local_filename, subs):
    merge_docx_header_footer(config, full_local_filename, subs, "header1")
    return merge_docx_header_footer(config, full_local_filename, subs, "header2")


def merge_docx_header_footer(config, full_local_filename, subs, xmlname):
    docx_filename = full_local_filename
    f = open(docx_filename, 'rb')
    zip = zipfile.ZipFile(f)
    xml_content = zip.read('word/'+xmlname+'.xml')

    xml_content = xml_content.decode("UTF-8")
    try:
        xml_content = clean_tags_in_word(xml_content)
        xml_content = substituteVariablesPlainString(config, xml_content, subs)
    except:
        pass
    tmp_dir = tempfile.mkdtemp()
    zip.extractall(tmp_dir)

    with open(os.path.join(tmp_dir, 'word/'+xmlname+'.xml'), 'w') as f:
        f.write(xml_content)
    filenames = zip.namelist()
    zip_copy_filename = docx_filename
    with zipfile.ZipFile(zip_copy_filename, "w") as docx:
        for filename in filenames:
            docx.write(os.path.join(tmp_dir,filename), filename)
    shutil.rmtree(tmp_dir)
    return({"file": docx_filename})
"""

def docx_relationship_dict(zip, filename):
    xml_content = zip.read(filename)
    xml_content = xml_content.decode("UTF-8")
    xml_content = preprocess(xml_content)
    xml_content = xml_content.replace("&quot;", '"')
    return [dict(d) for d in xmltodict.parse(xml_content)['Relationships']['Relationship']]

def docx_subfile(config, zip, tmp_dir, subs, filename):
    try:
        xml_content = zip.read(filename)
        xml_content = xml_content.decode("UTF-8")
        xml_content = preprocess(xml_content)
        xml_content = xml_content.replace("&quot;", '"')
        try:
            xml_content_subs = str(substituteVariablesPlainString(config, xml_content, subs))
        except django.template.exceptions.TemplateSyntaxError:
            raise TemplateError("Error in Template Structure: "+filename)

        with io.open(os.path.join(tmp_dir,filename), 'w', encoding="UTF-8") as f:
            f.write(xml_content_subs)
    except KeyError:
        pass

def docx_xml_part(zip, part_filename):
    xml_content = zip.read(part_filename)
    xml_content = xml_content.decode("UTF-8")
    xml_content = preprocess(xml_content)
    xml_content = xml_content.replace("&quot;", '"')
    return xml_content


def build_keys_list(doc, prefix=None):
    img_files = []
    for key in doc.keys():
        node = doc[key]
        if isinstance(node, dict):
            img_files += build_keys_list(node, prefix = key)
        elif isinstance(node, list):
            for item in node:
                if isinstance(item, dict):
                    img_files += build_keys_list(item, prefix=key)
        else:
            if key.rfind(".") > 1:
                if prefix:
                    ext_key = ".".join([prefix, key])
                else:
                    ext_key = key
                img_files.append((ext_key, doc[key.replace(".","_")+"_file"]))
    return img_files


def docx_subfile_subst_images(config, zip, subs, tmp_dir, filename, part):
    image_copies = []
    try:
        with io.open(os.path.join(tmp_dir,filename), 'r', encoding="UTF-8") as f:
            xml_content = f.read()
        xml_content = preprocess(xml_content)
        xml_content = xml_content.replace("&quot;", '"')
        keys_list = build_keys_list(subs)
        for pair in keys_list:
            target = pair[0]
            sub = pair[1]
            start = 0
            found_target = False
            abandon = False
            maxit = 20
            it = 0
            while not found_target and it < maxit and not abandon:
                find_snippet = xml_content[start:].find(target)
                if find_snippet < 0:
                    break
                snippet = xml_content[start:][find_snippet-10:find_snippet+5000]
                if snippet.find('name="') >= 0:
                    name_start = snippet[snippet.find('name="')+6:]
                    name_snip = name_start[:name_start.find('"')]
                    if snippet.find('r:embed="') >= 0:
                        id_start = snippet[snippet.find('r:embed="')+9:]
                        id_snip = id_start[:id_start.find('"')]
                    else:
                        id_snip = ''
                    image_copies.append((sub, name_snip, id_snip, part))
                    found_target = True
                else:
                    start = start+find_snippet+1
                it += 1
        with io.open(os.path.join(tmp_dir, filename), 'w', encoding="UTF-8") as f:
            f.write(xml_content)
    except KeyError:
        pass
    return image_copies


def docx_content(docx_filename):
    f = open(docx_filename, 'rb')
    zip = zipfile.ZipFile(f)
    part_filenames = zip.namelist()
    tmp_dir = tempfile.mkdtemp()
    zip.extractall(tmp_dir)

    content = {}

    for part in part_filenames:
        if part[-4:] == '.xml':
            content[part] = docx_xml_part(zip, part)

    return content


def substituteVariablesDocx_direct(config, file_name_in, file_name_out, subs):
    docx_filename = file_name_in
    f = open(docx_filename, 'rb')
    zip = zipfile.ZipFile(f)
    filenames = zip.namelist()
    tmp_dir = tempfile.mkdtemp()
    zip.extractall(tmp_dir)

    relationships = {}

    parts = [
        "document.xml",
        "header1.xml",
        "header2.xml",
        "header3.xml",
        "footer1.xml",
        "footer2.xml",
        "footer3.xml"
    ]

    for part in parts:
        try:
            part_relationships = { 
                d['@Id']: d['@Target'] 
                for d in docx_relationship_dict(zip, "word/_rels/"+part+".rels")
            }
        except (ValueError, KeyError):
            part_relationships = {}
        relationships[part] = part_relationships

    image_subs = []
    for filename in filenames:
        if filename in ['word/'+part for part in parts]:
            docx_subfile(config, zip, tmp_dir, subs, filename)
            image_subs += docx_subfile_subst_images(config, zip, subs, tmp_dir, filename, filename.replace('word/',''))

    with zipfile.ZipFile(file_name_out, "a") as docx:
        written_files = []
        for image_sub in image_subs:
            # if relationship id exists, decode it
            if image_sub[2] in relationships[image_sub[3]]:
                substitute = "word/"+relationships[image_sub[3]][image_sub[2]]
            else:
                substitute = "word/media/"+image_sub[1]
            docx.write(image_sub[0], substitute)
            written_files.append(substitute)
        for filename in filenames:
            if filename not in written_files:
                docx.write(os.path.join(tmp_dir, filename), filename)
        docx.close()
    shutil.rmtree(tmp_dir)

    return({"file": file_name_out})
"""

def get_docx_paras(zip):
    xml_content = zip.read("word/document.xml").decode("utf8")
    paras_start = xml_content.find("<w:p")
    paras_end = xml_content.rfind("</w:p>")+6
    return xml_content[paras_start:paras_end]


def get_docx_numbering(zip):
    xml_content = zip.read("word/numbering.xml").decode("utf8")
    remain = xml_content
    abstract_num_schemes = {}
    num_schemes = {}
    while remain.find("<w:abstractNum ")>=0:
        scheme_start = remain.find("<w:abstractNum ")
        scheme_end = remain.find("</w:abstractNum>")+16
        scheme_str = remain[scheme_start:scheme_end]
        scheme_index_loc_s = scheme_str.find('="')
        scheme_index_loc_e = scheme_str.find('">')
        scheme_index = scheme_str[scheme_index_loc_s+2:scheme_index_loc_e]
        abstract_num_schemes[scheme_index] = scheme_str
        remain = remain[scheme_end:]
    while remain.find("<w:num ")>=0:
        scheme_start = remain.find("<w:num ")
        scheme_end = remain.find("</w:num>")+8
        scheme_str = remain[scheme_start:scheme_end]
        scheme_index_loc_s = scheme_str.find('="')
        scheme_index_loc_e = scheme_str.find('">')
        scheme_index = scheme_str[scheme_index_loc_s+2:scheme_index_loc_e]
        num_schemes[scheme_index] = scheme_str
        remain = remain[scheme_end:]
    return abstract_num_schemes, num_schemes, xml_content


def get_docx_content(filename):
    docx_filename = filename
    f = open(docx_filename, 'rb')
    zip = zipfile.ZipFile(f)
    filenames = zip.namelist()
    tmp_dir = tempfile.mkdtemp()
    zip.extractall(tmp_dir)
    xml_str = zip.read("word/_rels/document.xml.rels").decode("utf8")
    rel_xml_content = strip_xml_dec(xml_str)
    rel_dom = etree.fromstring(rel_xml_content)
    response = {}
    response["paras"] = get_docx_paras(zip)
    response["numbering"] = get_docx_numbering(zip)
    response["tmp_dir"] = tmp_dir
    response["filenames"] = filenames
    response["rel_dom"] = rel_dom
    return response

def copy_docx_media(tmp_dir_from, tmp_dir_to, filenames_from, filenames_to, rel_0, max_Rel_id):
    rel_elements = []
    renames = []
    for filename in filenames_from:
        if filename.find("word/media/") == 0:
            source = os.path.join(tmp_dir_from, filename)
            dest = os.path.join(tmp_dir_to, filename)
            if not os.path.exists(os.path.split(dest)[0]):
                os.makedirs(os.path.split(dest)[0])
            if filename in filenames_to:
                oldfilename = filename
                i = 1
                while filename in filenames_to:
                    filename = "word/media/image"+"{0:0>2}".format(i)+".png"
                    i += 1
                renames.append((os.path.split(oldfilename)[1], os.path.split(filename)[1]))
                dest = os.path.join(tmp_dir_to, filename)
            else:
                renames.append((os.path.split(filename)[1], os.path.split(filename)[1]))
            shutil.copyfile(source.replace("/", os.sep), dest.replace("/", os.sep))
            filenames_to.append(filename)
            new_rel_el = etree.Element(
                "Relationship",
                attrib={
                    "Id": "".join(["rId", str(max_Rel_id+1)]), 
                    "Type": rel_0.attrib["Type"].replace("settings", "image"), 
                    "Target": "/".join(["media", os.path.split(filename)[1]])
                }
            )
            rel_elements.append(new_rel_el)
            max_Rel_id += 1
    return filenames_to, rel_elements, renames


def combine_docx_direct(file_names_to_combine, file_name_out):
    docx_filename = file_names_to_combine[0]
    f = open(docx_filename, 'rb')
    zip = zipfile.ZipFile(f)
    filenames = zip.namelist()
    tmp_dir = tempfile.mkdtemp()
    zip.extractall(tmp_dir)
    main_xml_content = zip.read("word/document.xml").decode("utf8")
    xml_str = zip.read("word/_rels/document.xml.rels").decode("utf8")
    rel_xml_content = strip_xml_dec(xml_str)
    rel_xml_dec = get_xml_dec(xml_str)
    rel_dom = etree.fromstring(rel_xml_content)
    max_Rel_id = len(rel_dom)
    rel_0 = rel_dom[0]
    main_abs_num, main_num, main_number_xml = get_docx_numbering(zip)
    insertion_point = main_xml_content.find("<w:sectPr>")
    for file in file_names_to_combine[1:]:
        if file == "pagebreak":
            break_xml = '<w:p><w:r><w:br w:type="page"/></w:r></w:p>'
            main_xml_content = main_xml_content[:insertion_point]+break_xml+main_xml_content[insertion_point:]
            insertion_point += len(break_xml)
        else:    
            sub_xml_content = get_docx_content(file)
            sub_xml_paras = sub_xml_content["paras"]
            abs_num, num, sub_num_xml = sub_xml_content["numbering"]
            if not main_abs_num:  # No main scheme
                if abs_num:  # But there is sub-document numbering
                    main_number_xml = sub_num_xml
                    main_abs_num = abs_num
                    main_num = num
            else:  # there is a main scheme
                if abs_num:  # And there is sub-document numbering
                    for key in abs_num.keys():
                        n_main_schemes = len(main_abs_num)
                        new_key = str(n_main_schemes+1)
                        main_abs_num[new_key]=abs_num[key].replace('w:abstractNumId="'+key+'"', 'w:abstractNumId="'+new_key+'"')
                        main_num[new_key]=num[key].replace('w:numId="'+key+'"', 'w:numId="'+new_key+'"').replace('<w:abstractNumId w:val="'+key+'"/>','<w:abstractNumId w:val="'+new_key+'"/>')
                        sub_xml_paras = sub_xml_paras.replace('<w:numId w:val="'+key+'"/>', '<w:numId w:val="'+new_key+'"/>')

            # Now move media files:  
            filenames, rel_elements, renames = copy_docx_media(sub_xml_content["tmp_dir"], tmp_dir, sub_xml_content["filenames"], filenames, rel_0, max_Rel_id)

            for rename in renames:
                sub_xml_paras = sub_xml_paras.replace(rename[0], rename[1])
                for element in rel_elements:
                    if element.attrib["Target"] == "media/"+rename[1]:
                        new_rel = element.attrib["Id"]

                        # make sure the replacement is right

                for element in sub_xml_content["rel_dom"]:
                    if element.attrib["Target"] == "media/"+rename[0]:
                        old_rel = element.attrib["Id"]
                sub_xml_paras = sub_xml_paras.replace(old_rel,new_rel)
 
            for element in rel_elements:
                rel_dom.append(element)

            max_Rel_id += len(rel_elements)

            main_xml_content = main_xml_content[:insertion_point]+sub_xml_paras+main_xml_content[insertion_point:]
            insertion_point+=len(sub_xml_paras)
            if main_number_xml.find("<w:abstractNum ") >=0:
                main_number_xml_pre = main_number_xml[:main_number_xml.find("<w:abstractNum ")]
                main_number_xml_post = main_number_xml[main_number_xml.rfind("</w:num>")+8:]
                main_number_xml = main_number_xml_pre
                for scheme in main_abs_num.values():
                    main_number_xml+=scheme
                for scheme in main_num.values():
                    main_number_xml+=scheme
                main_number_xml+=main_number_xml_post
    rels_xml = "\n".join([rel_xml_dec,etree.tostring(rel_dom).decode("utf8")])

    # All files assimilated
    with io.open(os.path.join(tmp_dir, "word/document.xml"), 'w', encoding="utf8") as f:
        f.write(main_xml_content)
    with io.open(os.path.join(tmp_dir, "word/numbering.xml"), 'w', encoding="utf8") as f:
        f.write(main_number_xml)
    with io.open(os.path.join(tmp_dir, "word/_rels/document.xml.rels"), 'w', encoding="utf8") as f:
        f.write(rels_xml)
    
    # rewrite the rel file
    with zipfile.ZipFile(file_name_out, "w") as docx:
        for filename in filenames:
            docx.write(os.path.join(tmp_dir, filename), filename)
    shutil.rmtree(tmp_dir)
    return({"file": file_name_out})


def shellCommand(command):
    return check_output(command, shell=True).decode()


def convert_pdf(filename_in, filename_out, outdir="."):
    command = "soffice --headless --convert-to pdf "+filename_in+" --outdir "+outdir
    response = shellCommand(command)
    return {"file": filename_out, "response": response, "command": command}


def convert_pdf_abiword(filename_in, filename_out, outdir = "."):
    command = "abiword --to=pdf "+filename_in+" --to-name="+filename_out
    response = shellCommand(command)
    return {"file":filename_out, "response":response, "command": command}


def watermark_pdf(target, wmark):
    optarget = target.replace(".pdf",".wm.pdf")
    target_file = PyPDF2.PdfFileReader(open(target, "rb"))
    wmark_file = PyPDF2.PdfFileReader(open(wmark, "rb"))
    output_file = PyPDF2.PdfFileWriter()
    for i in range(0, target_file.numPages):
        pageObj = target_file.getPage(i)
        print ("Watermarking page {} of {}".format(i+1, target_file.numPages))
        pageObj.mergePage(wmark_file.getPage(0))
        output_file.addPage(pageObj)
    with open(optarget, "wb") as outputStream:
        output_file.write(outputStream)
    return {"filename":optarget}


def password_pdf(target, password):
    optarget = target.replace(".pdf",".pw.pdf")
    target_file = PyPDF2.PdfFileReader(open(target, "rb"))
    output_file = PyPDF2.PdfFileWriter()

    for i in range(0, target_file.getNumPages()):
        output_file.addPage(target_file.getPage(i))
 
    outputStream = open(optarget, "wb")
 
    # Set user and owner password to pdf file
    output_file.encrypt(password, password, use_128bit=True)
    output_file.write(outputStream)
    outputStream.close()

    return {"filename": optarget}

"""
def clean_tag(tag):
    intrusive_tag_start = tag.find("<")
    intrusive_tag_end = tag.find(">")
    while intrusive_tag_start >= 0 and intrusive_tag_end > intrusive_tag_start:
        tag = ''.join([tag[:intrusive_tag_start], tag[intrusive_tag_end+1:]])
        intrusive_tag_start = tag.find("<")
        intrusive_tag_end = tag.find(">")
    return tag.replace("‘", "'").replace("’", "'")


def clean_tags_in_word(text, tag_def):
    done = []
    remaining = text
    while tag_def[0] in remaining:
        tagspan = (remaining.find(tag_def[0]), remaining.find(tag_def[1])+2)
        tag0 = remaining[tagspan[0]: tagspan[1]]
        tag1 = clean_tag(tag0)
        done.append(remaining[:tagspan[0]])
        done.append(tag1)
        remaining = remaining[tagspan[1]:]
    done.append(remaining)
    return ''.join(done)


def preprocess(text):
    text = text.replace("‘","'",).replace("’","'",)
    text = text.replace("{% #A", "{% cycle 'A' 'B' 'C' 'D' 'E' 'F' 'G' 'H' 'I' 'J' 'K' 'L' 'M' 'N' 'O' 'P' 'Q' 'R' 'S' 'T' 'U' 'V' 'W' 'X' 'Y' 'Z'")    
    text = text.replace("{% #a", "{% cycle 'a' 'b' 'c' 'd' 'e' 'f' 'g' 'h' 'i' 'j' 'k' 'l' 'm' 'n' 'o' 'p' 'q' 'r' 's' 't' 'u' 'v' 'w' 'x' 'y' 'z'")    
    text = text.replace("{% #9", "{% cycle '1' '2' '3' '4' '5' '6' '7' '8' '9' '10' '11' '12' '13' '14' '15' '16' '17' '18' '19' '20' '21' '22' '23' '24' '25' '26' '27' '28' '29' '30' '31' '32' '33' '34' '35' '36' '37' '38' '39' '40' '41' '42' '43' '44' '45' '46' '47' '48' '49' '50' '51' '52' '53' '54' '55' '56' '57' '58' '59' '60' '61' '62' '63' '64' '65' '66' '67' '68' '69' '70' '71' '72' '73' '74' '75' '76' '77' '78' '79' '80' '81' '82' '83' '84' '85' '86' '87' '88' '89' '40' '91' '92' '93' '94' '95' '96' '97' '98' '99' '100' ")    
    text = text.replace("{% #I", "{% cycle 'I' 'II' 'III' 'IV' 'V' 'VI' 'VII' 'VIII' 'IX' 'X' 'XI' 'XII' 'XIII' 'XIV' 'XV' 'XVI' 'XVII' 'XVIII' 'XIX' 'XX'")    
    text = text.replace("{% #i", "{% cycle 'i' 'ii' 'iii' 'iv' 'v' 'vi' 'vii' 'viii' 'ix' 'x' 'xi' 'xii' 'xiii' 'xiv' 'xv' 'xvi' 'xvii' 'xviii' 'xix' 'xx'")    
    text = clean_tags_in_word(text, ('{{', '}}'))
    text = clean_tags_in_word(text, ('{%', '%}'))
    return text
